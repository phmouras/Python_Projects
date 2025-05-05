import os
import json
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import re
from datetime import datetime

class DocumentFillerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Preenchimento Automático de Documentos")
        self.root.geometry("640x480")
        
        # Definir ícone do aplicativo (se disponível)
        try:
            icon_path = os.path.join(os.getcwd(), "icon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception:
            pass  # Não fazer nada se falhar
        
        # Configurar tema
        style = ttk.Style()
        if 'clam' in style.theme_names():  # Preferir o tema 'clam' se disponível
            style.theme_use('clam')
        
        self.documents_info = {}
        self.selected_docs = []
        self.all_fields = {}
        self.field_values = {}
        
        # Tentar carregar diretório templates automaticamente
        self.docx_dir = self.find_templates_directory()
        
        self.setup_ui()
        self.load_json_models()

    def find_templates_directory(self):
        """Busca por uma pasta 'templates' no diretório atual ou em subdiretórios"""
        # Verificar se existe uma pasta 'templates' no diretório atual
        current_dir = os.getcwd()
        templates_dir = os.path.join(current_dir, "templates")
        
        if os.path.exists(templates_dir) and os.path.isdir(templates_dir):
            return templates_dir
        
        # Verificar em subdiretórios imediatos
        for item in os.listdir(current_dir):
            item_path = os.path.join(current_dir, item)
            if os.path.isdir(item_path):
                templates_subdir = os.path.join(item_path, "templates")
                if os.path.exists(templates_subdir) and os.path.isdir(templates_subdir):
                    return templates_subdir
        
        # Se não encontrar, retornar None
        return None

    def setup_ui(self):
        # Frame principal com notebooks para diferentes etapas
        self.main_notebook = ttk.Notebook(self.root)
        self.main_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Página 1: Seleção de documentos
        self.docs_frame = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.docs_frame, text="1. Seleção de Documentos")
        
        # Página 2: Preenchimento de dados
        self.data_frame = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.data_frame, text="2. Preenchimento de Dados")
        
        # Página 3: Geração de documentos
        self.gen_frame = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.gen_frame, text="3. Geração de Documentos")
        
        # Configurar a página de seleção de documentos
        self.setup_docs_page()
        
    def setup_docs_page(self):
        # Label de instrução
        ttk.Label(self.docs_frame, text="Selecione os documentos a serem preenchidos:", 
                 font=("Arial", 12)).pack(anchor=tk.W, padx=20, pady=(20, 10))
        
        # Frame para checkboxes de documentos
        self.docs_checkbox_frame = ttk.Frame(self.docs_frame)
        self.docs_checkbox_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Lista de modelos de documentos disponíveis
        self.doc_vars = {}
        self.doc_checkboxes = []
        
        # Se temos um diretório de templates, já mostrar os modelos
        if self.docx_dir:
            self.load_docx_from_dir(self.docx_dir)
        else:
            ttk.Label(self.docs_checkbox_frame, 
                     text="Nenhum diretório de templates encontrado automaticamente.\nSelecione o diretório manualmente.",
                     font=("Arial", 10)).pack(pady=10)
        
        # Botão para avançar
        ttk.Button(self.docs_frame, text="Avançar para Preenchimento", 
                  command=self.proceed_to_data).pack(pady=20)
        
        # Botão para selecionar diretório de modelos .docx
        ttk.Button(self.docs_frame, text="Selecionar Diretório de Modelos", 
                  command=self.select_docx_dir).pack(pady=10)
                  
    def select_docx_dir(self):
        dir_path = filedialog.askdirectory(title="Selecione o diretório contendo os modelos .docx")
        if dir_path:
            self.docx_dir = dir_path
            self.load_docx_from_dir(dir_path)
    
    def load_docx_from_dir(self, dir_path):
        """Carrega os arquivos .docx do diretório especificado"""
        # Limpar os checkboxes existentes
        for cb in self.doc_checkboxes:
            cb.destroy()
        self.doc_checkboxes = []
        self.doc_vars = {}
        
        # Procurar por arquivos .docx no diretório
        try:
            docx_files = [f for f in os.listdir(dir_path) if f.endswith('.docx')]
            
            if not docx_files:
                messagebox.showinfo("Informação", "Nenhum arquivo .docx encontrado no diretório selecionado.")
                return
                
            # Criar checkboxes para cada arquivo .docx
            for idx, docx_file in enumerate(docx_files):
                var = tk.BooleanVar()
                cb = ttk.Checkbutton(self.docs_checkbox_frame, text=docx_file, variable=var)
                cb.grid(row=idx // 2, column=idx % 2, sticky=tk.W, padx=10, pady=5)
                self.doc_vars[docx_file] = var
                self.doc_checkboxes.append(cb)
                
            # Mostrar o caminho do diretório
            ttk.Label(self.docs_checkbox_frame, 
                     text=f"Diretório de modelos: {dir_path}",
                     font=("Arial", 8)).grid(row=(len(docx_files) // 2) + 1, column=0, columnspan=2, 
                                           sticky=tk.W, padx=10, pady=5)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar arquivos do diretório: {str(e)}")
    
    def load_json_models(self):
        # Carregar informações de campos de todos os modelos JSON
        # 1. Verificar no diretório atual
        cwd = os.getcwd()
        config_path = os.path.join(cwd, "config")
        self.load_json_from_dir(config_path)
        
        # 2. Se temos um diretório de templates, verificar lá também
        if self.docx_dir and self.docx_dir != os.getcwd():
            self.load_json_from_dir(self.docx_dir)
    
    def load_json_from_dir(self, directory):
        """Carrega os modelos JSON de um diretório específico"""
        for json_file in os.listdir(directory):
            if json_file.endswith('_modelo.json'):
                try:
                    with open(os.path.join(directory, json_file), 'r', encoding='utf-8') as f:
                        json_data = json.load(f)
                        if 'campos' in json_data:
                            self.documents_info[json_file] = json_data
                except Exception as e:
                    print(f"Erro ao carregar {json_file}: {e}")
    
    def proceed_to_data(self):
        # Verificar quais documentos foram selecionados
        self.selected_docs = [doc for doc, var in self.doc_vars.items() if var.get()]
        
        if not self.selected_docs:
            messagebox.showwarning("Aviso", "Selecione pelo menos um documento para continuar.")
            return
        
        # Identificar todos os campos necessários dos documentos selecionados
        self.identify_required_fields()
        
        # Configurar a página de preenchimento de dados
        self.setup_data_page()
        
        # Mudar para a próxima aba
        self.main_notebook.select(1)
   
    def identify_required_fields(self):
        self.all_fields = {}
        missing_jsons = []
        
        for selected_doc in self.selected_docs:
            base_name = selected_doc.split('.')[0]
            json_file = f"{base_name}.json"
                       
            if json_file not in self.documents_info:
                missing_jsons.append(selected_doc)
                continue
                
            try:
                fields = self.documents_info[json_file]['campos']
                for field_key, field_info in fields.items():
                    self.all_fields[field_key] = field_info
            except KeyError as e:
                messagebox.showerror("Erro", f"Modelo JSON inválido para {json_file}: Campo 'campos' não encontrado")
                return
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao processar {json_file}: {str(e)}")
                return
        
        if missing_jsons:
            messagebox.showerror("Erro", 
                            f"Modelos JSON não encontrados para:\n- " + 
                            "\n- ".join(missing_jsons))
            return 
     
    def setup_data_page(self):
        # Limpar widgets anteriores
        for widget in self.data_frame.winfo_children():
            widget.destroy()
        
        # Criar canvas para scrollbar
        canvas = tk.Canvas(self.data_frame)
        scrollbar = ttk.Scrollbar(self.data_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Empacotar canvas e scrollbar
        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y")
        
        # Label de instrução
        ttk.Label(scrollable_frame, text="Preencha os campos para os documentos selecionados:", 
                 font=("Arial", 12)).grid(row=0, column=0, columnspan=2, sticky=tk.W, padx=20, pady=(20, 10))
        
        # Adicionar campos agrupados por tipo
        self.field_entries = {}
        row = 1
        
        for field_key, field_info in sorted(self.all_fields.items(), key=lambda x: x[1]['rotulo']):
            label_text = field_info['rotulo']
            if field_info.get('obrigatorio', False):
                label_text += " *"
            
            ttk.Label(scrollable_frame, text=label_text).grid(row=row, column=0, sticky=tk.W, padx=20, pady=5)
            
            field_type = field_info.get('tipo', 'radio')  # Garantir que temos um tipo, padrão é 'texto'
            
            if field_type == 'radio':
                # Para campos de radio buttons
                radio_frame = ttk.Frame(scrollable_frame)
                radio_var = tk.StringVar()
                
                # Verificar se existem opções definidas
                opcoes = field_info.get('opcoes', [])
                if not opcoes:
                    print(f"Aviso: Campo {field_key} do tipo radio não tem opções definidas")
                    opcoes = ["Opção 1", "Opção 2"]  # Opções padrão
                
                for i, opcao in enumerate(opcoes):
                    rb = ttk.Radiobutton(radio_frame, text=opcao, variable=radio_var, value=opcao)
                    rb.pack(side=tk.LEFT, padx=5)
                
                radio_frame.grid(row=row, column=1, sticky=tk.W, padx=5, pady=5)
                self.field_entries[field_key] = radio_var
                
                # Selecionar a primeira opção por padrão
                if opcoes:
                    radio_var.set(opcoes[0])

            elif field_type == 'data':
                # Para campos de data, usar um combobox para o dia, mês e ano
                date_frame = ttk.Frame(scrollable_frame)
                
                day_var = tk.StringVar()
                day_combo = ttk.Combobox(date_frame, width=3, textvariable=day_var)
                day_combo['values'] = [str(i).zfill(2) for i in range(1, 32)]
                day_combo.pack(side=tk.LEFT, padx=2)
                
                ttk.Label(date_frame, text="/").pack(side=tk.LEFT)
                
                month_var = tk.StringVar()
                month_combo = ttk.Combobox(date_frame, width=3, textvariable=month_var)
                month_combo['values'] = [str(i).zfill(2) for i in range(1, 13)]
                month_combo.pack(side=tk.LEFT, padx=2)
                
                ttk.Label(date_frame, text="/").pack(side=tk.LEFT)
                
                year_var = tk.StringVar()
                year_combo = ttk.Combobox(date_frame, width=5, textvariable=year_var)
                current_year = datetime.now().year
                year_combo['values'] = [str(i) for i in range(current_year-5, current_year+2)]
                year_combo.pack(side=tk.LEFT, padx=2)
                
                date_frame.grid(row=row, column=1, sticky=tk.W, padx=5, pady=5)
                self.field_entries[field_key] = (day_var, month_var, year_var)
            
            else:
                # Para campos de texto
                if "Resumo" in field_key or "abstract" in field_key:
                    # Para campos de resumo, usar Text
                    text_widget = tk.Text(scrollable_frame, height=5, width=50)
                    text_widget.grid(row=row, column=1, sticky=tk.W, padx=5, pady=5)
                    self.field_entries[field_key] = text_widget
                else:
                    # Para outros campos de texto
                    text_var = tk.StringVar()
                    ttk.Entry(scrollable_frame, textvariable=text_var, width=50).grid(
                        row=row, column=1, sticky=tk.W, padx=5, pady=5)
                    self.field_entries[field_key] = text_var
            
            row += 1
        
        # Botões para navegar
        btn_frame = ttk.Frame(scrollable_frame)
        btn_frame.grid(row=row, column=0, columnspan=2, pady=20)
        
        ttk.Button(btn_frame, text="Voltar", command=lambda: self.main_notebook.select(0)).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Avançar para Geração", command=self.proceed_to_generation).pack(side=tk.LEFT, padx=10)
        
        # Botão para carregar dados de JSON
        ttk.Button(btn_frame, text="Carregar Dados de JSON", command=self.load_data_from_json).pack(side=tk.LEFT, padx=10)
        
    def load_data_from_json(self):
        json_file = filedialog.askopenfilename(title="Selecione o arquivo JSON com os dados",
                                               filetypes=[("Arquivos JSON", "*.json")])
        if not json_file:
            return
            
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            # Preencher os campos com os dados do JSON
            for field_key, field_entry in self.field_entries.items():
                if field_key in data:
                    value = data[field_key]
                    
                    # Verificar o tipo de widget/variável
                    if isinstance(field_entry, tuple):  # Data
                        day_var, month_var, year_var = field_entry
                        if isinstance(value, str) and '/' in value:
                            day, month, year = value.split('/')
                            day_var.set(day)
                            month_var.set(month)
                            year_var.set(year)
                    elif isinstance(field_entry, tk.Text):  # Text widget
                        field_entry.delete('1.0', tk.END)
                        field_entry.insert('1.0', value)
                    else:  # StringVar (para texto ou radio)
                        field_entry.set(value)
                        
            messagebox.showinfo("Sucesso", "Dados carregados com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar dados: {str(e)}")
    
    def proceed_to_generation(self):
        # Validar campos obrigatórios
        missing_fields = []
        
        for field_key, field_info in self.all_fields.items():
            if field_info.get('obrigatorio', False):
                field_entry = self.field_entries.get(field_key)
                
                if isinstance(field_entry, tuple):  # Data
                    day_var, month_var, year_var = field_entry
                    if not (day_var.get() and month_var.get() and year_var.get()):
                        missing_fields.append(field_info['rotulo'])
                elif isinstance(field_entry, tk.Text):  # Text widget
                    if not field_entry.get('1.0', tk.END).strip():
                        missing_fields.append(field_info['rotulo'])
                elif isinstance(field_entry, tk.StringVar):  # StringVar
                    if not field_entry.get().strip():
                        missing_fields.append(field_info['rotulo'])
        
        if missing_fields:
            messagebox.showwarning("Campos obrigatórios", 
                                  f"Por favor, preencha os seguintes campos obrigatórios:\n- " + 
                                  "\n- ".join(missing_fields))
            return
        
        # Obter valores dos campos
        self.field_values = {}
        
        for field_key, field_entry in self.field_entries.items():
            if isinstance(field_entry, tuple):  # Data
                day_var, month_var, year_var = field_entry
                date_value = f"{day_var.get()}/{month_var.get()}/{year_var.get()}"
                self.field_values[field_key] = date_value
            elif isinstance(field_entry, tk.Text):  # Text widget
                text_value = field_entry.get('1.0', tk.END).strip()
                self.field_values[field_key] = text_value
            else:  # StringVar (para texto ou radio)
                self.field_values[field_key] = field_entry.get()
        
        # Configurar página de geração
        self.setup_gen_page()
        
        # Mudar para a próxima aba
        self.main_notebook.select(2)
    
    def setup_gen_page(self):
        # Limpar widgets anteriores
        for widget in self.gen_frame.winfo_children():
            widget.destroy()
        
        # Label de instrução
        ttk.Label(self.gen_frame, text="Revise os dados antes de gerar os documentos:", 
                 font=("Arial", 12)).pack(anchor=tk.W, padx=20, pady=(20, 10))
        
        # Frame para mostrar resumo dos dados
        summary_frame = ttk.Frame(self.gen_frame)
        summary_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Criar canvas com scrollbar
        canvas = tk.Canvas(summary_frame)
        scrollbar = ttk.Scrollbar(summary_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Empacotar canvas e scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Mostrar os documentos selecionados
        ttk.Label(scrollable_frame, text="Documentos a serem gerados:", 
                 font=("Arial", 10, "bold")).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))
        
        for i, doc in enumerate(self.selected_docs):
            ttk.Label(scrollable_frame, text=f"• {doc}").grid(row=i+1, column=0, columnspan=2, sticky=tk.W)
        
        # Mostrar os principais campos preenchidos
        row = len(self.selected_docs) + 2
        ttk.Label(scrollable_frame, text="Dados principais:", 
                 font=("Arial", 10, "bold")).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        row += 1
        
        # Mostrar todos os campos preenchidos em ordem alfabética do rótulo
        sorted_fields = sorted(self.all_fields.items(), key=lambda x: x[1]['rotulo'])
        for field_key, field_info in sorted_fields:
            if field_key in self.field_values:
                label_text = field_info['rotulo']
                ttk.Label(scrollable_frame, text=f"{label_text}:").grid(row=row, column=0, sticky=tk.W, padx=(10, 5))
                ttk.Label(scrollable_frame, text=self.field_values[field_key]).grid(row=row, column=1, sticky=tk.W)
                row += 1
        
        # Frame para botões
        btn_frame = ttk.Frame(self.gen_frame)
        btn_frame.pack(pady=20)
        
        ttk.Button(btn_frame, text="Voltar para Edição", 
                  command=lambda: self.main_notebook.select(1)).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Salvar Dados como JSON", 
                  command=self.save_data_to_json).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Gerar Documentos", 
                  command=self.generate_documents).pack(side=tk.LEFT, padx=10)
    
    def save_data_to_json(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".json", 
                                              filetypes=[("Arquivos JSON", "*.json")],
                                              title="Salvar dados como")
        if not file_path:
            return
            
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(self.field_values, f, ensure_ascii=False, indent=4)
            messagebox.showinfo("Sucesso", "Dados salvos com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar dados: {str(e)}")
    
    def generate_documents(self):
        # Pedir diretório para salvar os documentos gerados
        output_dir = filedialog.askdirectory(title="Selecione o diretório para salvar os documentos gerados")
        if not output_dir:
            return
        
        try:
            # Mostrar progresso
            progress_window = tk.Toplevel(self.root)
            progress_window.title("Gerando documentos...")
            progress_window.geometry("400x150")
            progress_window.transient(self.root)
            progress_window.grab_set()
            
            # Centralizar janela
            progress_window.update_idletasks()
            width = progress_window.winfo_width()
            height = progress_window.winfo_height()
            x = (self.root.winfo_screenwidth() // 2) - (width // 2)
            y = (self.root.winfo_screenheight() // 2) - (height // 2)
            progress_window.geometry(f"{width}x{height}+{x}+{y}")
            
            # Adicionar barra de progresso
            ttk.Label(progress_window, text="Gerando documentos, aguarde...").pack(pady=(20, 10))
            progress_bar = ttk.Progressbar(progress_window, orient="horizontal", length=300, mode="determinate")
            progress_bar.pack(pady=10)
            progress_label = ttk.Label(progress_window, text="Processando...")
            progress_label.pack(pady=10)
            
            progress_window.update()
            
            # Total de documentos para cálculo de progresso
            total_docs = len(self.selected_docs)
            docs_gerados = []
            
            # Processar cada documento selecionado
            for i, doc_name in enumerate(self.selected_docs):
                # Atualizar progresso
                progress_bar["value"] = (i / total_docs) * 100
                progress_label.config(text=f"Processando: {doc_name}")
                progress_window.update()
                
                doc_path = os.path.join(self.docx_dir, doc_name)
                
                # Carregar o documento
                doc = Document(doc_path)
                
                # Substituir os campos nos parágrafos
                for para in doc.paragraphs:
                    for field_key, field_value in self.field_values.items():
                        if field_key in para.text:
                            para.text = para.text.replace(field_key, field_value)
                
                # Substituir os campos nas tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for field_key, field_value in self.field_values.items():
                                    if field_key in paragraph.text:
                                        paragraph.text = paragraph.text.replace(field_key, field_value)
                
                # Substituir nos cabeçalhos e rodapés
                for section in doc.sections:
                    # Verificar se o cabeçalho/rodapé não é None antes de tentar iterar
                    if section.header and section.header.paragraphs:
                        for paragraph in section.header.paragraphs:
                            for field_key, field_value in self.field_values.items():
                                if field_key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(field_key, field_value)
                    
                    if section.footer and section.footer.paragraphs:
                        for paragraph in section.footer.paragraphs:
                            for field_key, field_value in self.field_values.items():
                                if field_key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(field_key, field_value)
                
                # Definir nome do arquivo de saída
                base_name = os.path.splitext(doc_name)[0]
                nome_aluno = self.field_values.get("[nome do aluno]", "").replace(" ", "_")
                if not nome_aluno:
                    nome_aluno = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                output_filename = f"{base_name}_{nome_aluno}.docx"
                output_path = os.path.join(output_dir, output_filename)
                
                # Salvar o documento
                doc.save(output_path)
                docs_gerados.append(output_filename)
            
            # Fechar janela de progresso
            progress_window.destroy()
            
            # Mostrar resultado
            result_window = tk.Toplevel(self.root)
            result_window.title("Documentos Gerados")
            result_window.geometry("500x400")
            result_window.transient(self.root)
            
            # Centralizar janela
            result_window.update_idletasks()
            width = result_window.winfo_width()
            height = result_window.winfo_height()
            x = (self.root.winfo_screenwidth() // 2) - (width // 2)
            y = (self.root.winfo_screenheight() // 2) - (height // 2)
            result_window.geometry(f"{width}x{height}+{x}+{y}")
            
            # Adicionar informações de sucesso
            ttk.Label(result_window, text="Documentos gerados com sucesso!", font=("Arial", 14, "bold")).pack(pady=(20, 10))
            ttk.Label(result_window, text=f"Diretório: {output_dir}").pack(pady=(0, 10))
            
            # Lista de documentos gerados com scrollbar
            frame = ttk.Frame(result_window)
            frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
            
            scrollbar = ttk.Scrollbar(frame)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            listbox = tk.Listbox(frame, yscrollcommand=scrollbar.set, font=("Arial", 10))
            for doc in docs_gerados:
                listbox.insert(tk.END, doc)
            
            listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.config(command=listbox.yview)
            
            # Botão para abrir o diretório
            def open_directory():
                if os.name == 'nt':  # Windows
                    os.startfile(output_dir)
                elif os.name == 'posix':  # macOS ou Linux
                    import subprocess
                    subprocess.Popen(['xdg-open', output_dir])
            
            ttk.Button(result_window, text="Abrir Diretório", command=open_directory).pack(pady=20)
            ttk.Button(result_window, text="Fechar", command=result_window.destroy).pack(pady=(0, 20))
        
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documentos: {str(e)}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documentos: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentFillerApp(root)
    root.mainloop()